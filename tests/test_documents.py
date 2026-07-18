from io import BytesIO

from docx import Document

from app.documents import chunk_markdown, detect_document_title, extract_text


def test_chunk_markdown_respects_limit_and_overlap():
    text = "甲" * 900
    chunks = chunk_markdown(text, max_chars=400, overlap=80)
    assert len(chunks) == 3
    assert all(0 < len(chunk) <= 400 for chunk in chunks)
    assert chunks[0][-80:] == chunks[1][:80]


def test_extract_utf8_text():
    content = "儿童乘车规则：购票时请使用有效身份证件。" * 2
    assert extract_text("rule.txt", content.encode("utf-8")) == content


def test_extract_docx_paragraph_and_table():
    document = Document()
    document.add_paragraph("学生票规则说明")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "适用对象"
    table.cell(0, 1).text = "符合条件的学生"
    output = BytesIO()
    document.save(output)

    text = extract_text("rule.docx", output.getvalue())
    assert "学生票规则说明" in text
    assert "适用对象 | 符合条件的学生" in text


def test_detect_title_prefers_yaml_metadata():
    text = "---\ntitle: 高铁儿童票规则\ncategory: child_ticket\n---\n\n正文内容"
    assert detect_document_title("rule.md", text.encode("utf-8"), text) == "高铁儿童票规则"


def test_detect_title_uses_heading_then_filename_fallback():
    assert detect_document_title("rule.txt", b"", "# 学生优惠票办理说明\n\n正文") == "学生优惠票办理说明"
    assert detect_document_title("flight_safety_notice.txt", b"", "https://example.com") == "flight safety notice"
