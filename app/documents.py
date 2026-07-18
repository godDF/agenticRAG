from __future__ import annotations

import asyncio
import hashlib
import io
import re
import uuid
from datetime import datetime
from pathlib import Path
from typing import TYPE_CHECKING, Any

import yaml
from docx import Document
from pypdf import PdfReader

from app.config import RAG_CATEGORIES, Settings
from app.chunk_identity import build_chunk_uid
from app.registry import DocumentRegistry
from app.retrieval import VectorStore

if TYPE_CHECKING:
    from app.ingestion import IngestionReporter


ALLOWED_EXTENSIONS = {".md", ".txt", ".pdf", ".docx"}


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError("文本文件编码无法识别，请使用 UTF-8")


def extract_text(filename: str, data: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix not in ALLOWED_EXTENSIONS:
        raise ValueError("仅支持 .md、.txt、.pdf、.docx 文件")
    if suffix in {".md", ".txt"}:
        text = _decode_text(data)
    elif suffix == ".pdf":
        reader = PdfReader(io.BytesIO(data))
        text = "\n\n".join((page.extract_text() or "").strip() for page in reader.pages)
        if not text.strip():
            raise ValueError("PDF 未提取到文本，可能是扫描件；当前版本暂不支持 OCR")
    else:
        document = Document(io.BytesIO(data))
        blocks = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                if any(cells):
                    blocks.append(" | ".join(cells))
        text = "\n\n".join(blocks)
    text = text.replace("\x00", "").strip()
    if len(text) < 20:
        raise ValueError("文档有效文本过少，无法建立知识索引")
    return text


def detect_document_title(filename: str, data: bytes, text: str) -> str:
    """Detect a human-readable title without using an LLM."""
    suffix = Path(filename).suffix.lower()

    if suffix == ".md" and text.lstrip().startswith("---"):
        try:
            _prefix, header, _body = text.lstrip().split("---", 2)
            title = str((yaml.safe_load(header) or {}).get("title") or "").strip()
            if title:
                return title[:200]
        except (ValueError, yaml.YAMLError):
            pass

    if suffix == ".docx":
        try:
            title = str(Document(io.BytesIO(data)).core_properties.title or "").strip()
            if title:
                return title[:200]
        except Exception:
            pass

    if suffix == ".pdf":
        try:
            metadata = PdfReader(io.BytesIO(data)).metadata
            title = str((metadata.title if metadata else "") or "").strip()
            if title and title.lower() not in {"untitled", "无标题"}:
                return title[:200]
        except Exception:
            pass

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line or line in {"---", "***"}:
            continue
        line = re.sub(r"^#{1,6}\s*", "", line)
        line = re.sub(r"^[-*+•]\s+", "", line).strip()
        if not line or re.match(r"^[A-Za-z_][\w-]*\s*:", line):
            continue
        if 2 <= len(line) <= 200 and not line.lower().startswith(("http://", "https://")):
            return line

    fallback = Path(filename).stem.strip().replace("_", " ").replace("-", " ")
    return (fallback or "未命名文档")[:200]


def chunk_markdown(text: str, max_chars: int = 400, overlap: int = 80) -> list[str]:
    paragraphs = [part.strip() for part in re.split(r"\n\s*\n", text) if part.strip()]
    chunks: list[str] = []
    current = ""
    for paragraph in paragraphs:
        if len(current) + len(paragraph) + 2 <= max_chars:
            current = f"{current}\n\n{paragraph}".strip()
            continue
        if current:
            chunks.append(current)
        if len(paragraph) <= max_chars:
            current = paragraph
            continue
        start = 0
        while start < len(paragraph):
            piece = paragraph[start : start + max_chars].strip()
            if piece:
                chunks.append(piece)
            start += max_chars - overlap
        current = ""
    if current:
        chunks.append(current)
    return chunks


def _frontmatter(metadata: dict[str, Any], text: str) -> str:
    header = yaml.safe_dump(metadata, allow_unicode=True, sort_keys=False).strip()
    return f"---\n{header}\n---\n\n{text.strip()}\n"


def read_normalized_markdown(path: Path) -> tuple[dict[str, Any], str]:
    raw = path.read_text(encoding="utf-8")
    if not raw.startswith("---"):
        raise ValueError("规范化文档缺少 YAML 元数据")
    _, header, body = raw.split("---", 2)
    return yaml.safe_load(header) or {}, body.strip()


class DocumentService:
    def __init__(self, settings: Settings, registry: DocumentRegistry, vectors: VectorStore):
        self.settings = settings
        self.registry = registry
        self.vectors = vectors

    async def upload(
        self,
        *,
        filename: str,
        data: bytes,
        title: str | None,
        category: str,
        source_name: str | None,
        source_url: str,
        updated_at: str,
        reporter: "IngestionReporter | None" = None,
    ) -> dict[str, Any]:
        if category not in RAG_CATEGORIES:
            raise ValueError("category 必须是六类知识之一")
        if not data:
            raise ValueError("上传文件为空")
        if len(data) > self.settings.max_upload_mb * 1024 * 1024:
            raise ValueError(f"文件不能超过 {self.settings.max_upload_mb} MB")
        event_id = None
        if reporter:
            event_id = await reporter.start(
                kind="tool",
                name="document_parser",
                label="解析上传文档",
                progress=10,
                details={
                    "filename": Path(filename).name,
                    "extension": Path(filename).suffix.lower(),
                    "size_bytes": len(data),
                },
            )
        text = await asyncio.to_thread(extract_text, filename, data)
        if reporter and event_id:
            await reporter.finish(
                event_id,
                progress=25,
                details={"extracted_characters": len(text)},
            )
        content_hash = hashlib.sha256(text.encode("utf-8")).hexdigest()
        detected_title = (title or "").strip() or await asyncio.to_thread(
            detect_document_title, filename, data, text
        )
        normalized_source_name = (source_name or "").strip() or "用户上传文档"
        document_id = uuid.uuid4().hex
        created_at = datetime.now().isoformat(timespec="seconds")
        normalized_path = self.settings.kb_dir / category / f"{document_id}.md"
        event_id = None
        if reporter:
            event_id = await reporter.start(
                kind="tool",
                name="markdown_normalizer",
                label="规范化 Markdown 与元数据",
                progress=27,
                details={"category": category},
            )
        metadata = {
            "title": detected_title,
            "category": category,
            "source_name": normalized_source_name,
            "source_url": source_url.strip(),
            "updated_at": updated_at.strip(),
            "document_id": document_id,
            "content_hash": content_hash,
        }
        normalized_markdown = _frontmatter(metadata, text)
        await asyncio.to_thread(normalized_path.write_text, normalized_markdown, encoding="utf-8")
        if reporter and event_id:
            await reporter.finish(
                event_id,
                progress=36,
                details={
                    "normalized_path": normalized_path.relative_to(self.settings.data_dir).as_posix(),
                    "content_hash": content_hash[:12],
                },
            )

        event_id = None
        if reporter:
            event_id = await reporter.start(
                kind="tool",
                name="markdown_chunker",
                label="切分知识片段",
                progress=38,
                details={"max_chars": 400, "overlap_chars": 80},
            )
        chunks = self._payload_chunks(metadata, text, normalized_path)
        if reporter and event_id:
            await reporter.finish(
                event_id,
                progress=48,
                details={
                    "chunk_count": len(chunks),
                    "average_chunk_chars": round(
                        sum(len(chunk["content"]) for chunk in chunks) / len(chunks)
                    ) if chunks else 0,
                },
            )
        try:
            chunk_count = await self.vectors.upsert_document(
                document_id, chunks, reporter=reporter
            )
            record = {
                **metadata,
                "original_filename": Path(filename).name,
                "normalized_path": normalized_path.relative_to(self.settings.data_dir).as_posix(),
                "chunk_count": chunk_count,
                "status": "indexed",
                "created_at": created_at,
            }
            event_id = None
            if reporter:
                event_id = await reporter.start(
                    kind="tool",
                    name="sqlite_registry",
                    label="登记文档与索引信息",
                    progress=95,
                    details={"document_id": document_id},
                )
            saved = await asyncio.to_thread(self.registry.add, record)
            if reporter and event_id:
                await reporter.finish(event_id, progress=99)
            return saved
        except Exception:
            await self.vectors.delete_document(document_id)
            normalized_path.unlink(missing_ok=True)
            raise

    def _payload_chunks(
        self, metadata: dict[str, Any], text: str, normalized_path: Path
    ) -> list[dict[str, Any]]:
        relative = normalized_path.relative_to(self.settings.data_dir).as_posix()
        payloads = []
        for index, content in enumerate(chunk_markdown(text)):
            payloads.append(
                {
                    **metadata,
                    "content": content,
                    "file_path": relative,
                    "chunk_index": index,
                    "chunk_uid": build_chunk_uid(
                        category=metadata.get("category"),
                        title=metadata.get("title"),
                        source_url=metadata.get("source_url"),
                        content=content,
                    ),
                    "untrusted_content": True,
                }
            )
        return payloads

    async def list(self) -> list[dict[str, Any]]:
        return await asyncio.to_thread(self.registry.list)

    async def delete(self, document_id: str) -> dict[str, Any]:
        record = await asyncio.to_thread(self.registry.get, document_id)
        await self.vectors.delete_document(document_id)
        path = self.settings.data_dir / record["normalized_path"]
        path.unlink(missing_ok=True)
        return await asyncio.to_thread(self.registry.delete, document_id)

    async def reindex(self, document_id: str) -> dict[str, Any]:
        record = await asyncio.to_thread(self.registry.get, document_id)
        path = self.settings.data_dir / record["normalized_path"]
        metadata, text = await asyncio.to_thread(read_normalized_markdown, path)
        chunks = self._payload_chunks(metadata, text, path)
        await self.vectors.delete_document(document_id)
        count = await self.vectors.upsert_document(document_id, chunks)
        await asyncio.to_thread(self.registry.update_index, document_id, count)
        return await asyncio.to_thread(self.registry.get, document_id)
